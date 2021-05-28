import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from selenium import webdriver
import time
from gne import GeneralNewsExtractor
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import requests
import json

targetDir = ''

document = Document()


def write_paragraph(paragraph):
    paragraph = document.add_paragraph(paragraph)
    # print(len(document.paragraphs))
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.space_after = Pt(12)


def write_picture(index):
    p_img = document.add_paragraph()
    p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_run_img = p_img.add_run()
    img_obj = p_run_img.add_picture(targetDir + '/' + index + ".jpeg", width=Inches(5))


def write_title(title):
    document.add_heading(title, level=1)

    # records = (
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam')
    # )
    #
    # table = document.add_table(rows=1, cols=3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    # for qty, id, desc in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = id
    #     row_cells[2].text = desc

    document.save('demo.docx')


# 保存图片到本地


def save_file_to_local(dir_name, urls):
    i = 1
    global targetDir
    for each_url in urls:

        # 保存图片的位置
        targetDir = os.path.join(os.path.dirname(os.path.abspath(__file__)), dir_name + '/images')
        # 不存在创建文件夹
        if not os.path.isdir(targetDir):
            os.makedirs(targetDir)
        # 跟据文章的图片格式进行处理
        r_pic = requests.get(each_url)
        # 创建指定目录
        t = os.path.join(targetDir, str(i) + '.jpeg')
        print('该文章共需处理' + str(len(urls)) + '张图片，正在处理第' + str(i) + '张……')
        # 指定绝对路径
        fw = open(t, 'wb')
        # 保存图片到本地指定目录
        fw.write(r_pic.content)
        i += 1
        # 将旧的链接或相对链接修改为直接访问本地图片
        # update_file(each_url, t, htmlDir, file_name)
        fw.close()


def main():
    driver = webdriver.Chrome('/home/mi/tools/chromedriver')
    driver.get(
        'https://mp.weixin.qq.com/s?__biz=MzUyNDM4NDc4NA==&mid=2247501581&idx=1&sn=f26c6238cae565804512fe637f55fca5&chksm=fa2ca872cd5b2164800820778d22305eb77bb91f2213cd9f7ecc97e33818a8ea9b07f01ec158&scene=21#wechat_redirect')
    time.sleep(3)
    extractor = GeneralNewsExtractor()
    result = extractor.extract(driver.page_source, with_body_html=True)
    print(result)
    #  json.dumps(): 对数据进行编码。
    # json.loads(): 对数据进行解码。
    res = json.dumps(result)
    data = json.loads(res)
    title = data['title']
    content = data['content']
    images = data['images']
    print("images = : " + str(len(images)))
    for image in images:
        print(image)
    content_list = content.split('\n')
    insert_factor = len(content_list) // len(images)
    print("insert_factor = : " + str(insert_factor))
    print("content_list = : " + str(len(content_list)))
    write_title(title)
    dir_name = title.replace(' ', '')
    save_file_to_local(dir_name, images)
    for content in content_list:
        index = content_list.index(content)
        picture_index = index // insert_factor
        if index > 0 and index % insert_factor == 0 and picture_index < len(images):
            write_picture(str(picture_index))
        print(content)
        write_paragraph(content)

    driver.quit()
    document.save('demo.docx')


if __name__ == '__main__':
    main()
